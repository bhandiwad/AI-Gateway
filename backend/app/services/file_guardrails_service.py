"""
File Guardrails Service

Extracts text content from uploaded files and applies guardrail processors.
Supports: Images (OCR), PDFs, DOCX, TXT, CSV, JSON files.
"""
import base64
import io
import re
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass
import structlog

logger = structlog.get_logger()


@dataclass
class FileGuardrailResult:
    """Result of applying guardrails to a file."""
    passed: bool
    message: str
    action: str
    file_type: str
    extracted_text: Optional[str] = None
    detected_issues: Optional[List[Dict[str, Any]]] = None
    redacted_content: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None


def extract_text_from_file(
    file_content: bytes,
    file_type: str,
    filename: Optional[str] = None
) -> Tuple[str, str]:
    """
    Extract text content from a file.
    
    Args:
        file_content: Raw file bytes
        file_type: MIME type or extension
        filename: Optional filename for type detection
        
    Returns:
        Tuple of (extracted_text, detected_file_type)
    """
    file_type_lower = file_type.lower()
    
    # Detect type from filename if needed
    if filename:
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        if ext in ['pdf']:
            file_type_lower = 'application/pdf'
        elif ext in ['docx']:
            file_type_lower = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif ext in ['doc']:
            file_type_lower = 'application/msword'
        elif ext in ['txt']:
            file_type_lower = 'text/plain'
        elif ext in ['csv']:
            file_type_lower = 'text/csv'
        elif ext in ['json']:
            file_type_lower = 'application/json'
        elif ext in ['png', 'jpg', 'jpeg', 'gif', 'webp']:
            file_type_lower = f'image/{ext}'
    
    try:
        # Plain text files
        if 'text/plain' in file_type_lower or file_type_lower.endswith('.txt'):
            return file_content.decode('utf-8', errors='ignore'), 'text/plain'
        
        # CSV files
        if 'text/csv' in file_type_lower or 'csv' in file_type_lower:
            return file_content.decode('utf-8', errors='ignore'), 'text/csv'
        
        # JSON files
        if 'application/json' in file_type_lower or 'json' in file_type_lower:
            return file_content.decode('utf-8', errors='ignore'), 'application/json'
        
        # PDF files
        if 'application/pdf' in file_type_lower or 'pdf' in file_type_lower:
            return _extract_text_from_pdf(file_content), 'application/pdf'
        
        # DOCX files
        if 'wordprocessingml' in file_type_lower or 'docx' in file_type_lower:
            return _extract_text_from_docx(file_content), 'application/docx'
        
        # Image files - OCR
        if file_type_lower.startswith('image/') or any(x in file_type_lower for x in ['png', 'jpg', 'jpeg', 'gif', 'webp']):
            return _extract_text_from_image(file_content), 'image'
        
        # Fallback - try to decode as text
        try:
            return file_content.decode('utf-8', errors='ignore'), 'unknown'
        except:
            return "", 'binary'
            
    except Exception as e:
        logger.warning("file_text_extraction_failed", error=str(e), file_type=file_type)
        return "", 'error'


def _extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF using PyPDF2 or pdfplumber."""
    try:
        # Try PyPDF2 first
        try:
            from PyPDF2 import PdfReader
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n".join(text_parts)
        except ImportError:
            pass
        
        # Try pdfplumber as fallback
        try:
            import pdfplumber
            pdf_file = io.BytesIO(content)
            text_parts = []
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n".join(text_parts)
        except ImportError:
            pass
        
        logger.warning("pdf_extraction_no_library", message="Install PyPDF2 or pdfplumber for PDF support")
        return "[PDF content - install PyPDF2 for extraction]"
        
    except Exception as e:
        logger.warning("pdf_extraction_failed", error=str(e))
        return ""


def _extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc_file = io.BytesIO(content)
        doc = Document(doc_file)
        text_parts = []
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_parts.append(cell.text)
        return "\n".join(text_parts)
    except ImportError:
        logger.warning("docx_extraction_no_library", message="Install python-docx for DOCX support")
        return "[DOCX content - install python-docx for extraction]"
    except Exception as e:
        logger.warning("docx_extraction_failed", error=str(e))
        return ""


def _extract_text_from_image(content: bytes) -> str:
    """Extract text from image using OCR (pytesseract or easyocr)."""
    try:
        # Try pytesseract first
        try:
            import pytesseract
            from PIL import Image
            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image)
            return text
        except ImportError:
            pass
        
        # Try easyocr as fallback
        try:
            import easyocr
            reader = easyocr.Reader(['en'])
            results = reader.readtext(content)
            return " ".join([result[1] for result in results])
        except ImportError:
            pass
        
        logger.warning("ocr_no_library", message="Install pytesseract or easyocr for image OCR")
        return "[Image content - install pytesseract for OCR]"
        
    except Exception as e:
        logger.warning("ocr_extraction_failed", error=str(e))
        return ""


def apply_file_guardrails(
    file_content: bytes,
    file_type: str,
    profile: Any,
    tenant_id: int,
    filename: Optional[str] = None
) -> FileGuardrailResult:
    """
    Apply guardrail processors to a file's content.
    
    Args:
        file_content: Raw file bytes
        file_type: MIME type
        profile: GuardrailProfile to apply
        tenant_id: Tenant ID
        filename: Optional filename
        
    Returns:
        FileGuardrailResult with pass/fail and any detected issues
    """
    from backend.app.services.profile_guardrails_service import apply_profile_guardrails
    
    # Extract text from file
    extracted_text, detected_type = extract_text_from_file(file_content, file_type, filename)
    
    if not extracted_text:
        return FileGuardrailResult(
            passed=True,
            message="No text content to analyze",
            action="allow",
            file_type=detected_type,
            extracted_text=""
        )
    
    # Create a synthetic message for guardrail processing
    synthetic_message = {
        "role": "user",
        "content": f"[FILE CONTENT - {filename or 'uploaded file'}]\n{extracted_text}"
    }
    
    # Apply profile guardrails
    result = apply_profile_guardrails(
        profile=profile,
        messages=[synthetic_message],
        stage="request",
        tenant_id=tenant_id
    )
    
    if not result.passed:
        return FileGuardrailResult(
            passed=False,
            message=result.message,
            action=result.action,
            file_type=detected_type,
            extracted_text=extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text,
            detected_issues=[{
                "processor": result.triggered_processor,
                "action": result.action,
                "message": result.message
            }],
            metadata=result.metadata
        )
    
    # Check if content was redacted
    redacted_content = None
    if result.processed_messages and result.processed_messages[0].get("content") != synthetic_message["content"]:
        redacted_content = result.processed_messages[0].get("content", "").replace(
            f"[FILE CONTENT - {filename or 'uploaded file'}]\n", ""
        )
    
    return FileGuardrailResult(
        passed=True,
        message="File passed all guardrail checks",
        action="allow",
        file_type=detected_type,
        extracted_text=extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text,
        redacted_content=redacted_content
    )


def scan_base64_content(
    base64_data: str,
    content_type: str,
    profile: Any,
    tenant_id: int
) -> FileGuardrailResult:
    """
    Scan base64-encoded file content for guardrail violations.
    
    Args:
        base64_data: Base64-encoded file content
        content_type: MIME type of the content
        profile: GuardrailProfile to apply
        tenant_id: Tenant ID
        
    Returns:
        FileGuardrailResult
    """
    try:
        # Decode base64
        file_content = base64.b64decode(base64_data)
        return apply_file_guardrails(
            file_content=file_content,
            file_type=content_type,
            profile=profile,
            tenant_id=tenant_id
        )
    except Exception as e:
        logger.warning("base64_decode_failed", error=str(e))
        return FileGuardrailResult(
            passed=True,
            message="Could not decode base64 content",
            action="allow",
            file_type="unknown"
        )


def scan_image_url(
    image_url: str,
    profile: Any,
    tenant_id: int
) -> FileGuardrailResult:
    """
    Download and scan an image URL for guardrail violations.
    
    Args:
        image_url: URL of the image
        profile: GuardrailProfile to apply
        tenant_id: Tenant ID
        
    Returns:
        FileGuardrailResult
    """
    try:
        import httpx
        
        # Download image
        response = httpx.get(image_url, timeout=10.0)
        response.raise_for_status()
        
        content_type = response.headers.get('content-type', 'image/png')
        
        return apply_file_guardrails(
            file_content=response.content,
            file_type=content_type,
            profile=profile,
            tenant_id=tenant_id
        )
    except Exception as e:
        logger.warning("image_url_download_failed", error=str(e), url=image_url[:100])
        return FileGuardrailResult(
            passed=True,
            message="Could not download image for scanning",
            action="allow",
            file_type="image"
        )


def extract_files_from_messages(messages: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Extract all file references from messages (images, base64 data).
    
    Returns list of dicts with: type, data, url, content_type
    """
    files = []
    
    for msg in messages:
        content = msg.get("content")
        if isinstance(content, list):
            for part in content:
                if isinstance(part, dict):
                    if part.get("type") == "image_url":
                        image_url = part.get("image_url", {})
                        url = image_url.get("url", "") if isinstance(image_url, dict) else image_url
                        
                        if url.startswith("data:"):
                            # Base64 data URL
                            match = re.match(r'data:([^;]+);base64,(.+)', url)
                            if match:
                                files.append({
                                    "type": "base64",
                                    "content_type": match.group(1),
                                    "data": match.group(2)
                                })
                        else:
                            # External URL
                            files.append({
                                "type": "url",
                                "url": url,
                                "content_type": "image"
                            })
    
    return files


def apply_guardrails_to_message_files(
    messages: List[Dict[str, Any]],
    profile: Any,
    tenant_id: int,
    scan_images: bool = True
) -> Tuple[bool, str, List[Dict[str, Any]]]:
    """
    Scan all files in messages for guardrail violations.
    
    Args:
        messages: Chat messages that may contain files
        profile: GuardrailProfile to apply
        tenant_id: Tenant ID
        scan_images: Whether to OCR scan images
        
    Returns:
        Tuple of (passed, message, list of scan results)
    """
    files = extract_files_from_messages(messages)
    
    if not files:
        return True, "No files to scan", []
    
    results = []
    all_passed = True
    failure_message = ""
    
    for file_info in files:
        if file_info["type"] == "base64":
            result = scan_base64_content(
                base64_data=file_info["data"],
                content_type=file_info["content_type"],
                profile=profile,
                tenant_id=tenant_id
            )
        elif file_info["type"] == "url" and scan_images:
            result = scan_image_url(
                image_url=file_info["url"],
                profile=profile,
                tenant_id=tenant_id
            )
        else:
            continue
        
        results.append({
            "file_type": result.file_type,
            "passed": result.passed,
            "message": result.message,
            "issues": result.detected_issues
        })
        
        if not result.passed:
            all_passed = False
            failure_message = result.message
            break
    
    return all_passed, failure_message or "All files passed guardrail checks", results
